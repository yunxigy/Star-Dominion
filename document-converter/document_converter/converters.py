from __future__ import annotations

import base64
import html as html_lib
import importlib.util
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Iterable, Sequence
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Font


class ConversionError(RuntimeError):
    """A user-facing conversion error with a safe explanation."""


SUPPORTED_TARGETS = {
    "pdf-to-word-image",
    "office-to-pdf",
    "pdf-table-to-xlsx",
    "markdown-to-docx",
    "html-to-docx",
    "scan-to-docx",
}

TARGET_EXTENSIONS: dict[str, set[str]] = {
    "pdf-to-word-image": {".pdf"},
    "office-to-pdf": {".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".odt", ".ods", ".odp"},
    "pdf-table-to-xlsx": {".pdf"},
    "markdown-to-docx": {".md", ".markdown"},
    "html-to-docx": {".html", ".htm"},
    "scan-to-docx": {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".pdf"},
}

MAX_FILE_SIZE = 50 * 1024 * 1024
MAX_BATCH_FILES = 20
MAX_BATCH_SIZE = 200 * 1024 * 1024


def _module_available(name: str) -> bool:
    return importlib.util.find_spec(name) is not None


def _soffice() -> str | None:
    configured = os.environ.get("LIBREOFFICE_BIN") or os.environ.get("SOFFICE_PATH")
    if configured:
        configured_path = Path(configured).expanduser()
        if configured_path.is_file():
            return str(configured_path)
        resolved = shutil.which(configured)
        if resolved:
            return resolved
    return shutil.which("soffice") or shutil.which("libreoffice")


def dependency_capabilities() -> dict[str, bool]:
    return {
        "libreoffice": _soffice() is not None,
        "pdf": _module_available("fitz"),
        "pdf_tables": _module_available("pdfplumber") and _module_available("openpyxl"),
        "ocr": _module_available("pytesseract") and shutil.which("tesseract") is not None,
    }


def validate_input_for_target(path: Path, target: str) -> None:
    if target not in SUPPORTED_TARGETS:
        raise ConversionError(f"不支持的转换类型：{target}")
    if path.suffix.lower() not in TARGET_EXTENSIONS[target]:
        extensions = ", ".join(sorted(TARGET_EXTENSIONS[target]))
        raise ConversionError(f"不支持此文件类型；当前转换只支持 {extensions}，收到 {path.suffix or '无扩展名'}")


def _require_capability(name: str, message: str) -> None:
    if not dependency_capabilities()[name]:
        raise ConversionError(message)


def _safe_stem(path: Path) -> str:
    stem = re.sub(r"[^\w\-.一-龥 ]+", "_", path.stem).strip(" ._")
    return stem[:80] or "converted-document"


def _make_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = document.styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal.font.size = Pt(10.5)
    return document


def _text(node) -> str:
    return " ".join(str(node.get_text(" ", strip=True)).split())


def _add_inline_runs(paragraph, node) -> None:
    for child in node.children:
        if getattr(child, "name", None) == "br":
            paragraph.add_run().add_break()
        elif getattr(child, "name", None) in {"strong", "b", "em", "i", "code"}:
            run = paragraph.add_run(_text(child))
            run.bold = child.name in {"strong", "b"}
            run.italic = child.name in {"em", "i"}
        elif getattr(child, "name", None) is None:
            value = str(child)
            if value.strip():
                paragraph.add_run(value)
        else:
            _add_inline_runs(paragraph, child)


def _write_image(doc: Document, src: str, base_dir: Path) -> None:
    image_path: Path | None = None
    if src.startswith("data:"):
        match = re.match(r"data:([^;]+);base64,(.+)", src, re.DOTALL)
        if match:
            extension = mimetypes.guess_extension(match.group(1)) or ".bin"
            fd, temp_name = tempfile.mkstemp(suffix=extension)
            os.close(fd)
            image_path = Path(temp_name)
            image_path.write_bytes(base64.b64decode(match.group(2)))
    else:
        candidate = (base_dir / src).resolve()
        if candidate.is_file() and candidate.is_relative_to(base_dir.resolve()):
            image_path = candidate
    if image_path:
        try:
            doc.add_picture(str(image_path), width=Inches(6.2))
        finally:
            if src.startswith("data:"):
                image_path.unlink(missing_ok=True)


def _append_html_node(doc: Document, node, base_dir: Path, list_style: str | None = None) -> None:
    name = getattr(node, "name", None)
    if not name:
        return
    if name in {"script", "style", "noscript", "template"}:
        return
    if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        doc.add_heading(_text(node), level=int(name[1]))
        return
    if name in {"p", "blockquote", "div", "section", "article", "header", "footer"}:
        paragraph = doc.add_paragraph(style="Intense Quote" if name == "blockquote" else None)
        _add_inline_runs(paragraph, node)
        for image in node.find_all("img", recursive=True):
            _write_image(doc, image.get("src", ""), base_dir)
        if not paragraph.text.strip() and not node.find("img"):
            paragraph._element.getparent().remove(paragraph._element)
        return
    if name in {"ul", "ol"}:
        style = "List Bullet" if name == "ul" else "List Number"
        for child in node.find_all("li", recursive=False):
            _append_html_node(doc, child, base_dir, style)
        return
    if name == "pre":
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(node.get_text("", strip=False))
        run.font.name = "Courier New"
        return
    if name == "li":
        paragraph = doc.add_paragraph(style=list_style or "List Bullet")
        _add_inline_runs(paragraph, node)
        return
    if name == "table":
        rows = node.find_all("tr", recursive=True)
        if not rows:
            return
        columns = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max(columns, 1))
        table.style = "Table Grid"
        for row_idx, row in enumerate(rows):
            cells = row.find_all(["th", "td"], recursive=False)
            for col_idx, cell in enumerate(cells):
                table.cell(row_idx, col_idx).text = _text(cell)
                if cell.name == "th":
                    for run in table.cell(row_idx, col_idx).paragraphs[0].runs:
                        run.bold = True
        return
    if name == "img":
        _write_image(doc, node.get("src", ""), base_dir)
        return
    for child in node.find_all(recursive=False):
        _append_html_node(doc, child, base_dir, list_style)


def convert_html_to_docx(source: Path, output: Path) -> Path:
    from bs4 import BeautifulSoup

    raw = source.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    for element in soup(["script", "style", "noscript", "template"]):
        element.decompose()
    document = _make_document()
    root = soup.body or soup
    for child in root.find_all(recursive=False):
        _append_html_node(document, child, source.parent)
    if not document.paragraphs and not document.tables:
        document.add_paragraph(_text(root))
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def convert_markdown_to_docx(source: Path, output: Path) -> Path:
    import markdown

    html = markdown.markdown(
        source.read_text(encoding="utf-8", errors="replace"),
        extensions=["tables", "fenced_code", "sane_lists"],
    )
    temporary_file = tempfile.NamedTemporaryFile(prefix="markdown-", suffix=".html", dir=source.parent, delete=False)
    temporary = Path(temporary_file.name)
    temporary_file.close()
    try:
        temporary.write_text(html, encoding="utf-8")
        return convert_html_to_docx(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)


def convert_office_to_pdf(source: Path, output_dir: Path) -> Path:
    executable = _soffice()
    if not executable:
        raise ConversionError("服务器未安装 LibreOffice，请先安装 libreoffice-core、writer、calc、impress")
    output_dir.mkdir(parents=True, exist_ok=True)
    profile_dir = Path(tempfile.mkdtemp(prefix="libreoffice-profile-"))
    try:
        command = [
            executable,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(source),
        ]
        result = subprocess.run(command, capture_output=True, text=True, timeout=300)
        output = output_dir / f"{source.stem}.pdf"
        if result.returncode != 0 or not output.exists():
            detail = (result.stderr or result.stdout).strip()[-500:]
            raise ConversionError(f"LibreOffice 转换失败{': ' + detail if detail else ''}")
        return output
    except subprocess.TimeoutExpired as exc:
        raise ConversionError("Office 转 PDF 超时，请缩小文件或稍后重试") from exc
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)


def _fitz_document(source: Path):
    _require_capability("pdf", "服务器缺少 PyMuPDF，暂时无法处理 PDF 页面")
    import fitz

    try:
        return fitz.open(source)
    except Exception as exc:
        raise ConversionError("PDF 文件无法打开或已损坏") from exc


def convert_pdf_to_word_images(source: Path, output: Path) -> Path:
    pdf = _fitz_document(source)
    import fitz

    document = Document()
    output.parent.mkdir(parents=True, exist_ok=True)
    try:
        for index, page in enumerate(pdf):
            rect = page.rect
            section = document.sections[0] if index == 0 else document.add_section(WD_SECTION.NEW_PAGE)
            section.page_width = Inches(max(rect.width / 72, 1))
            section.page_height = Inches(max(rect.height / 72, 1))
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_path = output.parent / f".{output.stem}-page-{index + 1}.png"
            pixmap.save(str(image_path))
            try:
                document.add_picture(str(image_path), width=section.page_width)
            finally:
                image_path.unlink(missing_ok=True)
        if not pdf.page_count:
            raise ConversionError("PDF 不包含任何页面")
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        return output
    finally:
        pdf.close()


def _normalise_table(table: Iterable[Iterable[object]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table:
        rows.append(["" if cell is None else str(cell).strip() for cell in row])
    width = max((len(row) for row in rows), default=0)
    return [row + [""] * (width - len(row)) for row in rows]


def convert_pdf_table_to_xlsx(source: Path, output: Path) -> Path:
    _require_capability("pdf_tables", "服务器缺少 pdfplumber，无法可靠抽取 PDF 表格")
    import pdfplumber

    workbook = Workbook()
    workbook.remove(workbook.active)
    table_count = 0
    try:
        with pdfplumber.open(source) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                for table_number, table in enumerate(page.extract_tables(), start=1):
                    rows = _normalise_table(table)
                    if not rows:
                        continue
                    sheet_name = f"Page{page_number}_Table{table_number}"[:31]
                    sheet = workbook.create_sheet(sheet_name)
                    for row in rows:
                        sheet.append(row)
                    for cell in sheet[1]:
                        cell.font = Font(bold=True)
                    sheet.freeze_panes = "A2"
                    table_count += 1
        if table_count == 0:
            raise ConversionError("未识别到表格；扫描 PDF 请使用“扫描件转可编辑 Word”")
        output.parent.mkdir(parents=True, exist_ok=True)
        workbook.save(output)
        return output
    finally:
        workbook.close()


def _ocr_languages(pytesseract) -> str:
    try:
        languages = set(pytesseract.get_languages(config=""))
        return "chi_sim+eng" if "chi_sim" in languages else "eng"
    except Exception:
        return "eng"


def convert_scan_to_docx(source: Path, output: Path) -> Path:
    _require_capability("ocr", "服务器未安装 Tesseract OCR，请安装 tesseract、tesseract-langpack-chi_sim 和 pytesseract")
    import pytesseract
    from PIL import Image

    document = _make_document()
    images: list[Path] = []
    temporary_dir = Path(tempfile.mkdtemp(prefix="scan-images-"))
    try:
        if source.suffix.lower() == ".pdf":
            _require_capability("pdf", "服务器缺少 PyMuPDF，暂时无法处理 PDF 扫描件")
            import fitz

            pdf = _fitz_document(source)
            try:
                for index, page in enumerate(pdf):
                    image_path = temporary_dir / f"page-{index + 1}.png"
                    page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False).save(str(image_path))
                    images.append(image_path)
            finally:
                pdf.close()
        else:
            images.append(source)
        for index, image_path in enumerate(images):
            with Image.open(image_path) as image:
                text = pytesseract.image_to_string(image, lang=_ocr_languages(pytesseract)).strip()
            document.add_heading(f"第 {index + 1} 页 OCR 结果", level=2)
            if text:
                for paragraph in re.split(r"\n\s*\n", text):
                    document.add_paragraph("\n".join(line.strip() for line in paragraph.splitlines()).strip())
            else:
                document.add_paragraph("本页未识别到文字")
            document.add_paragraph("原始页面图片：")
            document.add_picture(str(image_path), width=Inches(6.2))
            if index < len(images) - 1:
                document.add_page_break()
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        return output
    finally:
        shutil.rmtree(temporary_dir, ignore_errors=True)


def convert_file(source: Path, target: str, output_dir: Path) -> Path:
    validate_input_for_target(source, target)
    stem = _safe_stem(source)
    if target == "pdf-to-word-image":
        return convert_pdf_to_word_images(source, output_dir / f"{stem}.docx")
    if target == "office-to-pdf":
        return convert_office_to_pdf(source, output_dir)
    if target == "pdf-table-to-xlsx":
        return convert_pdf_table_to_xlsx(source, output_dir / f"{stem}.xlsx")
    if target == "markdown-to-docx":
        return convert_markdown_to_docx(source, output_dir / f"{stem}.docx")
    if target == "html-to-docx":
        return convert_html_to_docx(source, output_dir / f"{stem}.docx")
    if target == "scan-to-docx":
        return convert_scan_to_docx(source, output_dir / f"{stem}.docx")
    raise ConversionError(f"不支持的转换类型：{target}")


def convert_batch(sources: Sequence[Path], target: str, work_dir: Path) -> Path:
    if not sources:
        raise ConversionError("批量转换至少需要一个文件")
    if len(sources) > MAX_BATCH_FILES:
        raise ConversionError(f"批量转换最多支持 {MAX_BATCH_FILES} 个文件")
    total_size = sum(path.stat().st_size for path in sources)
    if total_size > MAX_BATCH_SIZE:
        raise ConversionError("批量文件总大小不能超过 200 MiB")
    output_dir = work_dir / "converted"
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest: list[dict[str, str]] = []
    outputs: list[Path] = []
    for source in sources:
        try:
            result = convert_file(source, target, output_dir)
            outputs.append(result)
            manifest.append({"source": source.name, "status": "success", "output": result.name})
        except (ConversionError, OSError, ValueError) as exc:
            manifest.append({"source": source.name, "status": "failed", "error": str(exc)})
    if not outputs:
        raise ConversionError("所有文件均转换失败：" + json.dumps(manifest, ensure_ascii=False))
    manifest_path = output_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    archive = work_dir / "document-conversion-results.zip"
    with ZipFile(archive, "w", ZIP_DEFLATED) as zip_file:
        for output in outputs:
            zip_file.write(output, output.name)
        zip_file.write(manifest_path, "manifest.json")
    return archive
